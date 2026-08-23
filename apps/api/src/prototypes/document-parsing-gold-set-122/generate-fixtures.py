"""Generate deterministic high-density fixtures for Issue #122's prototype."""
from __future__ import annotations

import hashlib, io, json, os, re, shutil, zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
import pypdf._encryption as pypdf_encryption
import pypdfium2 as pdfium
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.pagesizes import A4, letter, landscape
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parent
PDF, DOCX, GOLD, CACHE = ROOT / "output" / "pdf", ROOT / "output" / "documents", ROOT / "gold", ROOT / "sources" / "cache"
FRONTIERS, NIST = CACHE / "frontiers-fphys-2020-00452.pdf", CACHE / "NIST.AI.100-1.pdf"
FRONTIERS_HASH = "EFE8CCE4394D43894BAA98D1ABEF12A3EAB69CEEF46B537822F470696D287A12"
NIST_HASH = "7576EDB531D9848825814EE88E28B1795D3A84B435B4B797D3670EAFDC4A89F1"
MAX_BYTES, ZIP_TIME = 20 * 1024 * 1024, (2026, 8, 24, 0, 0, 0)
NIST_NATIVE_PAGES = (4, 17, 27, 29, 31, 34, 37, 40, 43, 46)
NIST_RASTER_PAGES = (28, 45)
NIST_OCR_REFERENCE_TEXT = {
    28: "Mechanisms are in place to inventory AI systems",
    45: "AI Risk Management and Human-AI Interaction",
}
NIST_CITATION = "Tabassi, E. (2023), Artificial Intelligence Risk Management Framework (AI RMF 1.0), NIST AI 100-1, National Institute of Standards and Technology, Gaithersburg, MD, https://doi.org/10.6028/NIST.AI.100-1."
NIST_COURTESY = "Republished courtesy of the National Institute of Standards and Technology."

def font_path(*names):
    for name in names:
        path = Path(os.environ.get("WINDIR", "C:/Windows")) / "Fonts" / name
        if path.exists(): return str(path)
    raise FileNotFoundError(names)
def register_fonts():
    pdfmetrics.registerFont(TTFont("SC", font_path("NotoSansSC-VF.ttf", "msyh.ttc", "simhei.ttf")))
    pdfmetrics.registerFont(TTFont("SCB", font_path("NotoSansSC-VF.ttf", "msyhbd.ttc", "simhei.ttf")))
def sha(path): return hashlib.sha256(path.read_bytes()).hexdigest().upper()
def write_deterministic_encrypted_fixture(source, destination):
    """Create a stable benchmark fixture; these deterministic bytes are not cryptographic entropy."""
    writer=PdfWriter(); writer.append_pages_from_reader(PdfReader(source)); counter=0
    original_token_bytes=pypdf_encryption.secrets.token_bytes
    def deterministic_token_bytes(length):
        nonlocal counter
        result=bytearray()
        while len(result)<length:
            result.extend(hashlib.sha256(f"collector-122-encryption:{counter}".encode("ascii")).digest()); counter+=1
        return bytes(result[:length])
    try:
        pypdf_encryption.secrets.token_bytes=deterministic_token_bytes
        writer.encrypt("collector-122",algorithm="AES-256-R5")
        with destination.open("wb") as output: writer.write(output)
    finally:
        pypdf_encryption.secrets.token_bytes=original_token_bytes
def require_source(path, expected_hash, expected_bytes, expected_pages):
    if not path.exists() or path.stat().st_size != expected_bytes or sha(path) != expected_hash or len(PdfReader(path).pages) != expected_pages:
        raise RuntimeError(f"Invalid cached source {path}; run download-sources.py")
def new_canvas(path, size=A4):
    path.parent.mkdir(parents=True, exist_ok=True); c = canvas.Canvas(str(path), pagesize=size, pageCompression=1, invariant=1); c.setAuthor("Collector Prototype #122"); c.setCreator("deterministic fixture generator"); return c
def wrap(c, text, x, y, width, font="SC", size=10, leading=16):
    c.setFont(font, size); line = ""
    for char in text:
        trial = line + char
        if line and pdfmetrics.stringWidth(trial, font, size) > width: c.drawString(x, y, line); y -= leading; line = char
        else: line = trial
    if line: c.drawString(x, y, line); y -= leading
    return y
def enwrap(c, text, x, y, width, size=10, leading=14):
    c.setFont("Helvetica", size); line = ""
    for word in text.split():
        trial = (line + " " + word).strip()
        if line and pdfmetrics.stringWidth(trial, "Helvetica", size) > width: c.drawString(x, y, line); y -= leading; line = word
        else: line = trial
    if line: c.drawString(x, y, line); y -= leading
    return y
def chart_png():
    image = Image.new("RGB", (960, 480), "white"); draw = ImageDraw.Draw(image); regular = ImageFont.truetype(font_path("arial.ttf"), 25); bold = ImageFont.truetype(font_path("arialbd.ttf"), 29)
    draw.text((65, 22), "Synthetic evidence-recovery profile", font=bold, fill="#17324D"); draw.line((90,405,890,405), fill="#344054", width=4); draw.line((90,70,90,405), fill="#344054", width=4)
    for x, top, label in ((180,215,"Text"),(350,155,"Layout"),(520,102,"Source"),(690,187,"Degrade")): draw.rectangle((x,top,x+105,405), fill="#4F86C6", outline="#17324D", width=3); draw.text((x,428), label, font=regular, fill="#1D2939")
    out = io.BytesIO(); image.save(out, format="PNG"); return out.getvalue()
def chinese_furniture(c, page, chapter):
    width, height = A4; c.setStrokeColor(HexColor("#98A2B3")); c.line(54,height-45,width-54,height-45); c.setFillColor(HexColor("#475467")); c.setFont("SC",8); c.drawString(54,height-35,"证据链与可复核阅读 · #122 原创教材式样本"); c.drawRightString(width-54,height-35,chapter); c.line(54,63,width-54,63); c.setFont("SC",7.5); c.drawString(54,47,"页底注：匿名原创基准素材；页眉、页脚和图像均不得被猜写成正文。"); c.drawRightString(width-54,31,str(page)); c.setFillColor(black)
def draw_table(c, x, y, widths, rows, row_height=30, font="SC", size=8):
    for ri, row in enumerate(rows):
        xx=x
        for ci, value in enumerate(row):
            c.setFillColor(HexColor("#DCEAF7") if ri==0 else white); c.setStrokeColor(HexColor("#667085")); c.rect(xx,y-row_height,widths[ci],row_height,fill=1,stroke=1); c.setFillColor(black); c.setFont("SCB" if ri==0 else font,size); c.drawString(xx+5,y-19,value); xx+=widths[ci]
        y-=row_height
    return y

CHINESE_PAGES = [
 ("第一章 从阅读到证据链","1.1 研究型阅读的最小闭环","研究型阅读的最低目标不是把页面变成一段可复制文字，而是让读者能从一个判断回到原文。这个闭环至少包含连续正文、标题路径、页码和对象身份；其中任一环节断裂，摘要即使通顺，也不能用于复核。","例如，研究者看到“风险被优先处理”时，应能知道它来自哪一页、哪个标题、哪张表或哪段解释。系统若只输出句子，就把事实与证据位置拆开了。","本章采用“内容—结构—定位”三层检查：先确认原文是否完整，再确认阅读顺序，最后确认对象关系。三层结果分别记录，避免高可读性掩盖低可复核性。"),
 ("第一章 从阅读到证据链","1.2 标题层级与阅读顺序","标题不是装饰。一级标题给出主题范围，二级标题组织论证顺序，三级标题则把一个可操作的问题放进稳定位置。解析时保留层级，可以让使用者回到“为什么讨论这一点”的上下文。","阅读顺序同样不能靠页面坐标猜测。正文、提示框、页脚和图注可能在视觉上相邻，却承担不同角色。可靠结果要保留这种角色差异，而不是把它们按行拼接。","练习：比较“标题→正文→提示”与“正文→提示→标题”两种输出。前者可导航，后者虽然包含相同词语，却会让后续引用失去章节路径。"),
 ("第一章 从阅读到证据链","1.3 示例：同一句话的两种返回路径","同一句话可以有粗略路径，也可以有精确路径。粗略路径只给出页码；精确路径还给出标题、区域和对象类型。研究笔记不必总是精确到坐标，但应说明自身使用的是哪一级定位。","当句子来自表格单元格时，页码并不足够。读者还需要表名、列标题和行标识；当句子来自图像时，至少应保留图号、图注和图片本身是否已经可见。","因此，系统应把“可回到来源”视为质量属性，而不是导出后的可选附加项。后续章节会用跨页表、图和公式分别演示这种要求。"),
 ("第二章 结构对象不能退化为纯文本","2.1 跨页表格（上）","表格把多个判断压缩在行列关系中。将单元格按视觉顺序抄成句子，会丢失指标属于哪一列、例外属于哪一行等关键信息。本页给出表 1 的前半部分，下一页保留重复表头并继续。","表 1 的列名本身是解释的一部分：对象说明解析的目标，最低保留要求说明成功标准，错误降级的后果说明为什么不能只保留文本。","读者在引用本页时，应同时记录“表 1”“第 4 页”“上半段”和相关列名；只记录数字或句子会把原本的比较关系抹平。"),
 ("第二章 结构对象不能退化为纯文本","2.1 跨页表格（续）","跨页续表最容易暴露不稳定的提取流程：如果第二页缺少表名或重复表头，读者无法判断这些行是否仍属于同一张表。这里将表名写为“表 1（续）”，并保留列标题作为阅读锚点。","续页增加脚注、标题和例外项，目的是让候选方案同时保留正文引用和页底条件。表格对象应以行、列、合并单元格与续页关系被识别，而非一串无归属的短语。","完成本节后，检查输出能否回答：该行接在什么表之后？它的列标题是什么？页底注是否被误认为正文？无法回答时应披露结构丢失。"),
 ("第二章 结构对象不能退化为纯文本","2.2 图像、图注与正文交叉引用","图像提供的不是可有可无的装饰，而是另一种证据载体。正文提到图 1 时，读者应能得到图像资源、图注和与正文的关系；只留下图注会让图中的趋势或分组不可复核。","本页的图 1 是合成柱状图，用来区分文字、版面、来源和降级四个维度。它的结论不应被改写成正文事实；系统只能说明图像存在、图注存在以及图像是否成功可见。","当图像不可提取时，正确结果是标记“图片可见性缺失”或给出 OCR/视觉分析的来源，而不是根据相邻段落编造柱高或标签。"),
 ("第三章 定位的表达能力","3.1 页码不是唯一定位","页码是最低成本的返回方式，却不能承担全部定位职责。同一页可能有正文、边栏、表格、页脚和图像；若没有区域或对象类型，页码只能告诉读者大概位置，无法减少歧义。","稳定定位可以写成“第 7 页／3.1／正文上半部”，也可以写成“第 7 页／表 1（续）／第二列”。两者精度不同，但都应如实表达，不应伪装成像素级坐标。","版本也属于定位的一部分。页面内容因重排或修订改变后，旧页码可能仍存在而含义已经变化。因此，引用应连同文档版本或哈希一起保存。"),
 ("第三章 定位的表达能力","3.2 块公式与变量关系","块公式通常承载一段推理的核心关系。若只把符号当作普通文本，容易遗漏上下标、分式、变量说明或公式与邻近段落之间的解释关系。","公式 1 将整体证据分数拆成文字完整性、定位质量和视觉对象保留度。它是本教材的合成示例，不是产品阈值；真正重要的是公式、变量释义和所在章节被作为一个整体返回。","行内公式也需要保留，但它们不能代替块公式的对象身份。解析结果应能区分“正文中出现 S”与“公式 1 定义 S”，从而让复核者回到正确位置。"),
 ("第三章 定位的表达能力","3.3 行内公式与例外","行内公式适合表达局部条件，例如当 T 小于阈值时，应提示文字层不足；当 V 缺失时，不能用平均分覆盖图片缺失。它们往往夹在句子中，因此阅读顺序错误会直接改变条件与结论的关系。","例外条件尤其容易在页脚、括号或脚注中出现。系统若能恢复正文却遗漏例外，应把结果标为部分可用，而不是将主结论升级成无条件事实。","实践中，定位记录可同时含页码、标题和对象；不需要强迫每个行内表达式都有坐标，但不能在没有来源时声称拥有精确来源。"),
 ("第四章 诚实降级","4.1 不可识别对象的处理","无法可靠读取并不等于任务失败，但输出必须把不可读取的边界说清楚。扫描页可能没有原生文字层，图像可能没有替代文本，损坏文件甚至不能安全打开；这些情况不能被伪装成空白内容。","诚实降级至少包含三件事：说明失败对象，说明已有内容是否仍可用，说明不能作出的承诺。例如“本页为图像，未取得原生文字层”比“页面没有内容”更准确。","这一原则也约束定位：没有稳定区域时可以只返回页码或对象存在性，却不能虚构行号、段落号或完整句子。"),
 ("第四章 诚实降级","4.2 参考文献与版本","参考文献既是来源对象，也是版本线索。书目信息、文内引用和附录版本说明应能彼此关联；若只保留作者姓名而丢失标题、年份或页码，读者无法判断是否回到了同一来源。","本教材使用匿名原创参考条目，避免把练习材料误写成真实研究结论。参考文献区仍然需要作为独立区域被识别，而不是与正文段落混合。","复核者应检查：版本信息是否在输出中可见？文内引文是否能指向条目？条目是否被错误拆成表格或页脚？这些问题决定引用能否重建。"),
 ("附录 A 复核清单","A.1 对象级检查","附录将前文的原则转换成可执行检查：先读标题与正文，再检查表格、图像、公式和脚注，最后验证每个主张是否能回到页码与对象。对象级检查不取代整体价值判断，两者必须并列记录。","对于跨页表，应确认续页、重复表头和行列关系；对于图，应确认图像资源与图注；对于公式，应确认表达式与变量说明；对于扫描页，应确认是否存在原生文字层。","完成检查后，报告应把成功、部分成功和失败分别写出。一个没有伪造内容、同时明确标注缺失的结果，通常比看似完整却无法回源的结果更有研究价值。"),
]
def chinese_pdf(path):
    c = new_canvas(path, A4); width,height=A4; margin=54; body=width-2*margin
    top=[["对象","最低保留要求","错误降级的后果"],["表格","行列与单元格关系","数字脱离指标"],["公式","表达式与变量释义","变量关系丢失"],["图片","资源、图注和页码","图证据被写成正文"]]
    nxt=[["对象","最低保留要求","错误降级的后果"],["脚注","正文引用与页脚内容","例外条件丢失"],["标题","层级与序列","导航路径失效"],["版本","来源标识与页码","引用无法复现"]]
    for page,(chapter,section,a,b,d) in enumerate(CHINESE_PAGES,1):
        chinese_furniture(c,page,chapter); y=height-86; c.setFont("SCB",16); c.drawString(margin,y,chapter); c.setFillColor(HexColor("#1F4D78")); c.setFont("SCB",12); c.drawString(margin,y-27,section); c.setFillColor(black); y=wrap(c,a,margin,y-57,body,size=10.2,leading=17); y=wrap(c,b,margin,y-8,body,size=10.2,leading=17); y=wrap(c,f"本页复核记录围绕“{section}”展开：阅读者应同时比较内容、结构和返回路径，确认当前结论是否仍保有上下文。若页面中的任一对象不能稳定识别，应保留它的存在、位置和不确定性，而不是用相邻文字填补空缺。",margin,y-8,body,size=10.2,leading=17)
        if page==4: c.setFont("SCB",10); c.drawString(margin,y-13,"表 1 结构对象的研究后果（上）"); y=draw_table(c,margin,y-29,[120,165,195],top); y=wrap(c,d,margin,y-16,body,size=9.6,leading=15); c.setFillColor(HexColor("#EEF5FB")); c.roundRect(margin,y-128,body,104,5,fill=1,stroke=0); c.setFillColor(HexColor("#17324D")); c.setFont("SCB",9.5); c.drawString(margin+11,y-48,"案例检查") ; wrap(c,"如果“数字脱离指标”被输出为单独句子，读者就无法判断它属于表格、公式还是图注。记录表名、列名和续页关系，才能让比较恢复为原来的比较。",margin+11,y-66,body-22,size=8.8,leading=13); c.setFillColor(black)
        elif page==5: c.setFont("SCB",10); c.drawString(margin,y-13,"表 1（续）结构对象的研究后果"); y=draw_table(c,margin,y-29,[120,165,195],nxt); y=wrap(c,d,margin,y-16,body,size=9.6,leading=15); c.setFillColor(HexColor("#EEF5FB")); c.roundRect(margin,y-128,body,104,5,fill=1,stroke=0); c.setFillColor(HexColor("#17324D")); c.setFont("SCB",9.5); c.drawString(margin+11,y-48,"续页检查") ; wrap(c,"重复表头不是重复内容：它让第二页的行继续拥有列含义。若续页被截为普通段落，结果必须披露行列关系和表名已经丢失。",margin+11,y-66,body-22,size=8.8,leading=13); c.setFillColor(black)
        elif page==6: y=wrap(c,d,margin,y-8,body,size=10.2,leading=17); c.drawImage(ImageReader(io.BytesIO(chart_png())),margin,190,width=body,height=225,preserveAspectRatio=True); c.setFont("SC",9); c.drawCentredString(width/2,176,"图 1  四类解析对象的合成恢复概况（见本页正文）")
        elif page==8:
            c.setFillColor(HexColor("#EEF5FB")); c.roundRect(margin,y-112,body,99,5,fill=1,stroke=0); c.setFillColor(HexColor("#17324D")); c.setFont("SCB",10); c.drawString(margin+12,y-36,"公式 1  区域加权证据分数"); c.setFont("Helvetica",17); c.drawString(margin+42,y-66,"S = 0.45T + 0.35L + 0.20V"); c.setFont("SC",8.8); c.drawString(margin+42,y-88,"其中 T 为文字完整性，L 为定位质量，V 为视觉对象保留度。"); c.setFillColor(black); y=wrap(c,d,margin,y-137,body,size=10.2,leading=17); c.setFillColor(HexColor("#F6F8FA")); c.roundRect(margin,y-143,body,124,5,fill=1,stroke=0); c.setFillColor(HexColor("#17324D")); c.setFont("SCB",9.5); c.drawString(margin+11,y-43,"变量复核") ; wrap(c,"T、L、V 的含义必须和公式同时返回。若公式字符可见但变量说明丢失，结果只能说明公式片段存在，不能断言已经恢复了可解释的计算关系。对分式、上下标或矩阵等更复杂对象，同样需要以对象级结果说明保留程度。",margin+11,y-61,body-22,size=8.8,leading=13); c.setFillColor(black)
        elif page==11:
            y=wrap(c,d,margin,y-8,body,size=10.2,leading=17); c.setFont("Helvetica",8.8)
            for i,ref in enumerate(("[1] Lin, Q. Evidence-aware reading systems. Synthetic Press, 2026.","[2] Zhao, M. Stable locators for research notes. Example Journal, 9(1), 11-29.","[3] Collector Project. Anonymous benchmark protocol #122, 2026.","[4] Chen, R. Object-aware document recovery. Demonstration Review, 4(2), 33-48.")): c.drawString(margin,y-18-i*22,ref)
            c.setFillColor(HexColor("#EEF5FB")); c.roundRect(margin,y-300,body,142,5,fill=1,stroke=0); c.setFillColor(HexColor("#17324D")); c.setFont("SCB",9.5); c.drawString(margin+11,y-192,"引文返回检查") ; wrap(c,"引用输出至少应保留条目编号、作者或机构、标题线索和版本信息。若只恢复了一个姓名，不能把它报告为完整参考文献；若无法确定文内引文对应哪一条，应保留这种不确定性。附录 A 将把这些检查与表格、图像和公式并列执行。",margin+11,y-210,body-22,size=8.8,leading=13); c.setFillColor(black)
        else:
            y=wrap(c,d,margin,y-8,body,size=10.2,leading=17); c.setFillColor(HexColor("#EEF5FB")); c.roundRect(margin,max(91,y-166),body,151,5,fill=1,stroke=0); c.setFillColor(HexColor("#17324D")); c.setFont("SCB",9.5); c.drawString(margin+11,max(219,y-53),"复核提示"); wrap(c,f"本页检查点与“{section}”相连：确认标题路径、正文顺序、页码和对象类型能够共同返回来源。若某类对象缺失，请把缺失写成可见结果，而不是补写为完整内容。对有例外条件的段落，还应确认条件没有被页脚、提示框或图注遮蔽。",margin+11,max(201,y-72),body-22,size=8.7,leading=13); c.setFillColor(black)
        c.showPage()
    c.save()
def raster_nist_page(source_page):
    source=pdfium.PdfDocument(str(NIST)); image=source[source_page-1].render(scale=2).to_pil().convert("RGB"); pixels=image.load()
    for y in range(17,image.height,181):
        for x in range((y*13)%37,image.width,223): r,g,b=pixels[x,y]; pixels[x,y]=(max(0,r-15),max(0,g-15),max(0,b-15))
    out=io.BytesIO(); image.save(out,format="PNG"); return out.getvalue()
def derivative_page(kind):
    out=io.BytesIO(); size=landscape(letter) if kind=="wide" else letter; c=canvas.Canvas(out,pagesize=size,pageCompression=1,invariant=1); width,height=size
    if kind=="cover":
        c.setFillColor(HexColor("#17324D")); c.rect(0,0,width,height,fill=1,stroke=0); c.setFillColor(white); c.setFont("Helvetica-Bold",11); c.drawCentredString(width/2,height-190,"COLLECTOR ISSUE #122 · BENCHMARK DERIVATIVE"); c.setFont("Helvetica-Bold",27); c.drawCentredString(width/2,height-260,"Recoverable AI Risk Evidence"); c.setFont("Helvetica",15); c.drawCentredString(width/2,height-292,"A mixed digital / scanned technical-report stress sample"); c.setFont("Helvetica",9); c.drawCentredString(width/2,155,"Derivative of NIST AI 100-1 for parser benchmarking only"); c.drawCentredString(width/2,137,"NOT AN OFFICIAL NIST PUBLICATION · no endorsement implied")
    else:
        c.setStrokeColor(HexColor("#98A2B3")); c.line(50,height-42,width-50,height-42); c.setFillColor(HexColor("#475467")); c.setFont("Helvetica",8); c.drawString(50,height-32,"BENCHMARK DERIVATIVE · NOT AN OFFICIAL NIST PUBLICATION"); c.setFillColor(black); x,y,available=54,height-82,width-108
        if kind=="notice":
            c.setFont("Helvetica-Bold",18); c.drawString(x,y,"Source, license, and derivative notice")
            paras=[NIST_CITATION,NIST_COURTESY,"Modified on 2026-08-24 for parser evaluation: selected and reordered source pages; rasterized source PDF physical pages 28 and 45 with synthetic scan noise; added four original benchmark framing pages.","Source page numbers in this derivative are 1-based physical PDF indexes; the printed labels inside the source pages are recorded separately in the Gold set.","This benchmark derivative selects and rasterizes only the NIST pages listed in its page map. It avoids pages carrying third-party copyright notices, including text or figures explicitly credited to OECD.","This is not an official NIST publication. NIST does not review, approve, or endorse this derivative, its benchmark framing, or any result produced from it.","Output pages 6 and 12 intentionally contain no native PDF text layer. They are rasterized derivatives with light synthetic scan noise; OCR, if attempted, must be reported as OCR rather than as native text."]
            for p in paras: y=enwrap(c,p,x,y-31,available,size=10,leading=15)
        elif kind=="wide":
            c.setFont("Helvetica-Bold",18); c.drawString(x,y,"Wide control comparison (landscape)")
            rows=[["Control","Source return","Object relationship","Scan downgrade","Risk if lost","Reviewer evidence"],["Page number","Coarse","None","No","Ambiguous claim","Page + region"],["Heading path","Medium","Section only","No","Wrong context","Heading hierarchy"],["Table cell","Fine","Row / column","Partial","Metric drift","Header + cell"],["Raster page","Visual","Image only","Yes","False text","OCR or explicit absence"],["Source hash","Version","Document","No","Wrong edition","SHA-256 record"]]; widths=[105,105,140,110,150,138]
            for ri,row in enumerate(rows):
                row_y=y-35-ri*46; xx=x
                for ci,value in enumerate(row): c.setFillColor(HexColor("#DCEAF7") if ri==0 else white); c.setStrokeColor(HexColor("#667085")); c.rect(xx,row_y-36,widths[ci],36,fill=1,stroke=1); c.setFillColor(black); c.setFont("Helvetica-Bold" if ri==0 else "Helvetica",7.2); enwrap(c,value,xx+4,row_y-12,widths[ci]-8,size=7.2,leading=8.8); xx+=widths[ci]
            enwrap(c,"The table is original benchmark framing. Its purpose is to distinguish a page-level return path from a cell-level return path without copying an additional third-party table.",x,190,available,size=9.5)
        else:
            c.setFont("Helvetica-Bold",18); c.drawString(x,y,"Appendix C: provenance and page map")
            paras=[NIST_CITATION,NIST_COURTESY,"Cached source SHA-256: "+NIST_HASH+".","Modified on 2026-08-24 for parser evaluation: selected and reordered source pages; rasterized source PDF physical pages 28 and 45 with synthetic scan noise; added four original benchmark framing pages.","All source numbers below are 1-based physical PDF indexes, not the printed labels inside a page. Native output pages map to source PDF pages 4, 17, 27, 29, 31, 34, 37, 40, 43, and 46. Output pages 6 and 12 rasterize source PDF pages 28 and 45 respectively. The remaining output pages are original synthetic benchmark framing.","This derivative is not an official NIST publication and does not imply NIST endorsement. Source-page selection intentionally excludes pages marked with third-party copyright material."]
            for p in paras: y=enwrap(c,p,x,y-31,available,size=10,leading=15)
    c.showPage(); c.save(); return out.getvalue()
def image_only_page(image):
    out=io.BytesIO(); c=canvas.Canvas(out,pagesize=letter,pageCompression=1,invariant=1); c.drawImage(ImageReader(io.BytesIO(image)),0,0,width=letter[0],height=letter[1]); c.showPage(); c.save(); return out.getvalue()
def hybrid_pdf(path):
    source=PdfReader(NIST); writer=PdfWriter(); sequence=[("generated","cover"),("generated","notice"),("source",4),("source",17),("source",27),("raster",28),("source",29),("source",31),("generated","wide"),("source",34),("source",37),("raster",45),("source",40),("source",43),("source",46),("generated","provenance")]
    for mode,value in sequence:
        if mode=="source": writer.add_page(source.pages[value-1])
        elif mode=="raster": writer.add_page(PdfReader(io.BytesIO(image_only_page(raster_nist_page(value)))).pages[0])
        else: writer.add_page(PdfReader(io.BytesIO(derivative_page(value))).pages[0])
    path.parent.mkdir(parents=True,exist_ok=True)
    with path.open("wb") as f: writer.write(f)

def cell_margins(cell):
    props=cell._tc.get_or_add_tcPr(); margins=OxmlElement("w:tcMar"); props.append(margins)
    for name,value in (("top",45),("start",80),("bottom",45),("end",80)):
        element=OxmlElement("w:"+name); element.set(qn("w:w"),str(value)); element.set(qn("w:type"),"dxa"); margins.append(element)
def table_geometry(table,widths):
    table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.autofit=False; props=table._tbl.tblPr
    for name,value in (("tblW",sum(widths)),("tblInd",80)):
        existing=props.find(qn("w:"+name))
        if existing is not None: props.remove(existing)
        element=OxmlElement("w:"+name); element.set(qn("w:w"),str(value)); element.set(qn("w:type"),"dxa"); props.append(element)
    grid=table._tbl.tblGrid
    for column in list(grid): grid.remove(column)
    for width in widths:
        column=OxmlElement("w:gridCol"); column.set(qn("w:w"),str(width)); grid.append(column)
    for row in table.rows:
        for index,cell in enumerate(row.cells):
            cell_margins(cell); cell_props=cell._tc.get_or_add_tcPr(); existing=cell_props.find(qn("w:tcW"))
            if existing is not None: cell_props.remove(existing)
            element=OxmlElement("w:tcW"); element.set(qn("w:w"),str(widths[min(index,len(widths)-1)])); element.set(qn("w:type"),"dxa"); cell_props.append(element)
def shade(cell,color):
    element=OxmlElement("w:shd"); element.set(qn("w:fill"),color); cell._tc.get_or_add_tcPr().append(element)
def set_style(document,name,font,size,color,before,after,line):
    style=document.styles[name]; style.font.name=font; style.font.size=Pt(size); style.font.color.rgb=RGBColor.from_string(color); style._element.rPr.rFonts.set(qn("w:ascii"),font); style._element.rPr.rFonts.set(qn("w:hAnsi"),font); style.paragraph_format.space_before=Pt(before); style.paragraph_format.space_after=Pt(after); style.paragraph_format.line_spacing=line
def add_numpr(paragraph,level):
    props=paragraph._p.get_or_add_pPr(); number=OxmlElement("w:numPr"); level_node=OxmlElement("w:ilvl"); level_node.set(qn("w:val"),str(level)); number_id=OxmlElement("w:numId"); number_id.set(qn("w:val"),"5"); number.append(level_node); number.append(number_id); props.append(number)
def row_height(row,twips):
    props=row._tr.get_or_add_trPr(); height=OxmlElement("w:trHeight"); height.set(qn("w:val"),str(twips)); height.set(qn("w:hRule"),"atLeast"); props.append(height)
def mark_header_row(row):
    props=row._tr.get_or_add_trPr(); repeat=OxmlElement("w:tblHeader"); repeat.set(qn("w:val"),"true"); props.append(repeat)
def real_footnote(path):
    with zipfile.ZipFile(path) as z: files={name:z.read(name) for name in z.namelist()}
    files["word/document.xml"]=re.sub(br"<w:t[^>]*>\[\[FN\]\]</w:t>",b'</w:r><w:r><w:footnoteReference w:id="1"/></w:r><w:r>',files["word/document.xml"])
    files["word/footnotes.xml"]=b'<?xml version="1.0" encoding="UTF-8"?><w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote><w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote><w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r><w:r><w:t> This footnote is a real OOXML note used to test source-return handling.</w:t></w:r></w:p></w:footnote></w:footnotes>'
    files["word/_rels/document.xml.rels"]=files["word/_rels/document.xml.rels"].replace(b"</Relationships>",b'<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>')
    files["[Content_Types].xml"]=files["[Content_Types].xml"].replace(b"</Types>",b'<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>')
    with zipfile.ZipFile(path,"w",zipfile.ZIP_DEFLATED,compresslevel=9) as z:
        for name in sorted(files): info=zipfile.ZipInfo(name,ZIP_TIME); info.compress_type=zipfile.ZIP_DEFLATED; z.writestr(info,files[name])
def small_cell(cell,text):
    cell.text=text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after=Pt(0); paragraph.paragraph_format.line_spacing=.9
        for run in paragraph.runs: run.font.size=Pt(8)
def docx_fixture(path):
    d=Document(); section=d.sections[0]; section.page_width,section.page_height=Inches(8.5),Inches(11); section.top_margin=section.right_margin=section.bottom_margin=section.left_margin=Inches(1); section.header_distance=section.footer_distance=Inches(.492)
    for args in (("Normal","Calibri",10.5,"000000",0,5,1.08),("Heading 1","Calibri",16,"2E74B5",15,7,1),("Heading 2","Calibri",13,"2E74B5",11,5,1),("Heading 3","Calibri",11.5,"1F4D78",7,3,1)): set_style(d,*args)
    header=section.header.paragraphs[0]; header.text="RESEARCH OPERATIONS BRIEF  |  SYNTHETIC COMPLEX DOCX"; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; header.runs[0].font.size=Pt(8.5)
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run("Collector #122 · Page "); field=OxmlElement("w:fldSimple"); field.set(qn("w:instr")," PAGE "); footer._p.append(field)
    d.add_paragraph().paragraph_format.space_after=Pt(84); p=d.add_paragraph("BENCHMARK BUSINESS BRIEF"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.color.rgb=RGBColor(122,90,0); p.runs[0].font.size=Pt(10); p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Recoverable Evidence Operations"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor(32,55,72); p=d.add_paragraph("An eight-page office-document regression sample"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.size=Pt(14); d.add_paragraph("Original synthetic content · headings, lists, merged cells, tables, images, fields, and OOXML footnotes.").alignment=WD_ALIGN_PARAGRAPH.CENTER; d.add_page_break()
    d.add_paragraph("Executive summary",style="Heading 1"); d.add_paragraph("This document uses real Word styles and structures rather than visually imitated formatting. A useful parser should preserve headings, paragraph order, list levels, table cells, images, captions, and source-return notes."); d.add_paragraph("Decision criteria",style="Heading 2")
    for level,text in ((0,"Preserve semantic headings before styling details."),(0,"Keep lists as real numbered paragraphs, including nested levels."),(1,"Nested review: return each result to a source block."),(0,"Expose table merges and repeated headers without inventing cells."),(0,"Disclose missing images or footnotes explicitly.")):
        p=d.add_paragraph(text,style="List Number" if level==0 else "List Number 2"); p.paragraph_format.left_indent=Inches(.5+.25*level); p.paragraph_format.first_line_indent=Inches(-.25); add_numpr(p,level)
    d.add_paragraph("Source-return exception",style="Heading 3"); d.add_paragraph("[[FN]]"); d.add_page_break()
    d.add_paragraph("Observed results",style="Heading 1"); d.add_paragraph("Merged-cell decision matrix",style="Heading 2"); matrix=d.add_table(rows=4,cols=3); matrix.style="Table Grid"; table_geometry(matrix,[2100,3660,3600]); matrix.cell(0,0).merge(matrix.cell(0,2)); small_cell(matrix.cell(0,0),"Evidence recovery decision matrix"); shade(matrix.cell(0,0),"E8EEF5")
    for cell,value in zip(matrix.rows[1].cells,("Object","Recovery expectation","Honest downgrade")): small_cell(cell,value); shade(cell,"F2F4F7")
    mark_header_row(matrix.rows[0]); mark_header_row(matrix.rows[1])
    for row,values in zip(matrix.rows[2:],(("Headings","Style hierarchy and sequence","Keep raw text with hierarchy unknown"),("Footnotes","Reference and note body","Report note as unavailable"))):
        for cell,value in zip(row.cells,values): small_cell(cell,value); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    d.add_paragraph("Table 1. Merged title row and structured body cells.",style="Caption"); d.add_page_break()
    d.add_paragraph("Cross-page review register",style="Heading 1"); d.add_paragraph("The register contains 32 distinct records. Its first row is a repeating header and the row height makes the table continue from page 4 onto page 5 in Microsoft Word.")
    register=d.add_table(rows=1,cols=4); register.style="Table Grid"; table_geometry(register,[780,2100,2600,3480])
    for cell,value in zip(register.rows[0].cells,("ID","Object","Expected recovery","Failure disclosure")): small_cell(cell,value); shade(cell,"F2F4F7")
    mark_header_row(register.rows[0]); row_height(register.rows[0],360)
    objects=["Heading 1","Heading 2","Heading 3","Numbered list","Nested list","Merged cells","Table header","Table continuation","Inline image","Caption","Footnote","Page field","Header","Footer","Source list","Reference entry","Executive paragraph","Decision criterion","Exception note","Matrix cell","Register cell","Merged title","Alt text","Section break","Page break","Appendix heading","Appendix checklist","Provenance note","Version string","Object locator","Partial-result note","Final review"]
    for i,name in enumerate(objects,1):
        cells=register.add_row().cells
        for cell,value in zip(cells,(f"R-{i:02d}",name,f"Preserve {name.lower()} with a return path",f"If absent, disclose the {name.lower()} gap")): small_cell(cell,value); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row_height(register.rows[-1],720)
    d.add_page_break(); d.add_paragraph("Figure evidence",style="Heading 1"); d.add_paragraph("Inline image and caption",style="Heading 2"); temp=ROOT/"office-figure.tmp.png"; temp.write_bytes(chart_png()); inline=d.add_picture(str(temp),width=Inches(6.15)); inline._inline.docPr.set("descr","Bar chart of four synthetic evidence-recovery dimensions"); d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER; cap=d.add_paragraph("Figure 1. Synthetic evidence-recovery profile used only for image and caption extraction.",style="Caption"); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER; d.add_paragraph("Reproducibility notes",style="Heading 2"); d.add_paragraph("The chart is inline rather than floating so the caption has a stable neighbouring relationship. The register ends before this section; an extractor must not report the figure as an additional register row."); d.add_page_break()
    d.add_paragraph("Appendix A: review procedure",style="Heading 1"); d.add_paragraph("Procedure detail",style="Heading 2"); d.add_paragraph("Reviewers first confirm document-level usefulness, then check diagnostic anchors. The two levels are intentionally separate: a document can contain all words while still losing the relationships that make a research claim inspectable."); d.add_paragraph("Structure audit",style="Heading 3")
    for text in ("Confirm real Heading 1, Heading 2, and Heading 3 styles.","Confirm numbered paragraphs retain two levels.","Confirm merged table cell and repeated header are detectable.","Confirm inline image, Caption, footnote part, and PAGE field exist."): d.add_paragraph(text,style="List Bullet")
    d.add_paragraph("The long register on pages 4–5 is deliberately made of individual, meaningful records. A parser should preserve their order and header relationship, rather than joining them into one broad narrative paragraph."); d.add_page_break()
    d.add_paragraph("Appendix B: sources and final checklist",style="Heading 1"); d.add_paragraph("All prose, data, and the chart in this DOCX are original synthetic benchmark material. The appendix is retained on its own page to test long-document navigation as well as final-page extraction."); d.add_paragraph("Reference list",style="Heading 2"); d.add_paragraph("Collector Prototype. Complex office regression specification #122, 2026.",style="List Number"); d.add_paragraph("Example Standards Group. Recoverable document evidence. Synthetic edition, 2026.",style="List Number"); d.add_paragraph("Final review",style="Heading 2"); d.add_paragraph("Before accepting a result, confirm the source title, the repeated table header, the chart caption, the real footnote, and the footer page field. Any missing object must remain visible as a partial result rather than being filled with guessed content.")
    path.parent.mkdir(parents=True,exist_ok=True); d.save(path); temp.unlink(missing_ok=True); real_footnote(path)

def positive_checks(name):
    labels={
        "zh-single-column-textbook":["正文、章节标题、跨页表、图、公式、参考文献和附录完整，不靠重复段落制造页数。","章与节形成连续阅读路径；表 1 的上半段和续页仍是同一个对象。","图、公式、参考文献、页码与对象区域能支撑研究者回到原页。","无法提取某类对象时，明确对象类型与范围，不把视觉内容伪成正文。"],
        "en-two-column-paper":["Title、Abstract、正文、附录和 References 的范围完整，不跨栏拼句。","按视觉阅读顺序读取左右栏，图表跨栏时不打断章节逻辑。","图表、方程、附录表和文内引用均可回到页级对象。","无法恢复版式时明确标记双栏限制，不把交错文本报告为正常文章。"],
        "hybrid-technical-report":["封面、来源声明、目录、原生正文/表图、附录和页映射作为同一报告完整保留。","目录、章节标题与页映射让读者能跨原生页、横向页和扫描页理解报告结构。","原生图表、长表、横向宽表与扫描页分别说明自己的页级和对象级复核范围。","第 6、12 页没有原生文字层；OCR 结果必须标明来源，不能伪装成原生精确文字。"],
        "office-docx-regression":["段落、列表、表格、图片、脚注、页眉页码组成的内容完整。","Heading 1/2/3 和真实多级列表可用于顺序阅读。","图注、脚注、跨页表和页码能回到相应对象。","不支持嵌入对象时明确标明缺失，正文不被静默删去。"],
    }[name]
    return [{"id":identifier,"label":label} for identifier,label in zip(("integrity","navigable-reading","citation-return","honest-degradation"),labels)]
def negative_checks(name):
    labels={
        "negative-corrupted":["不凭空产生正文或摘要。","不把损坏字节伪装成可导航结构。","不产生看似可复核的页码或引用。","稳定返回 damaged_file / corrupt_pdf 专用失败分类。"],
        "negative-encrypted":["未提供密码时不返回正文。","不因元数据猜测出一份可导航文档。","不产生未经解密的可复核引用。","稳定返回 password_required / encrypted_pdf 专用失败分类。"],
        "negative-oversized":["大小校验失败后不产生正文。","不进入阅读结构或对象提取阶段。","不制造可复核的页内定位。","解析前返回 file_too_large，并报告 20 MiB 限制。"],
    }[name]
    return [{"id":identifier,"label":label} for identifier,label in zip(("integrity","navigable-reading","citation-return","honest-degradation"),labels)]
def sample(identifier,group,file,weight,layout,anchors,outcome="success",decision_dependency=None):
    result={"id":identifier,"group":group,"candidate":True,"file":file,"weight":weight,"expectedOutcome":outcome,"layout":layout,"documentValueChecks":negative_checks(identifier) if group=="negative" else positive_checks(identifier),"diagnosticAnchors":anchors}
    if decision_dependency is not None: result["decisionDependency"]=decision_dependency
    return result
def external_benchmarks():
    return [{
        "id":"omnidocbench-v1.6",
        "status":"approved-external-research-only",
        "includedInScore":False,
        "includedInRepository":False,
        "automaticDownload":False,
        "purpose":"Optional page-level comparison for text, table, formula, layout, and reading-order accuracy, with emphasis on the 296-page Hard subset.",
        "doesNotReplace":["whole-document research value","cross-page continuity","DOCX native structure","corrupt, encrypted, and oversized failure behavior"],
        "licenseBoundary":"Dataset copyright statement limits use to research and prohibits commercial use; no dataset asset or cache is committed or distributed.",
        "independenceBoundary":"OmniDocBench and MinerU are both in the OpenDataLab ecosystem, so this benchmark is not the sole evidence for choosing MinerU.",
        "revisionPolicy":"Any future local run must pin both the dataset revision and evaluation-code commit and report results separately from this seven-file score.",
        "references":{
            "dataset":"https://huggingface.co/datasets/opendatalab/OmniDocBench",
            "evaluationCode":"https://github.com/opendatalab/OmniDocBench",
            "copyrightStatement":"https://github.com/opendatalab/OmniDocBench#copyright-statement",
        },
    }]
def gold_set():
    return {"schemaVersion":2,"prototype":True,"issue":122,"decisionStatus":"awaiting-user-review","question":"Can a candidate preserve research value and disclose loss, rather than merely export text?","scoring":{"candidate":True,"weights":{"documentValue":60,"diagnosticAnchors":40},"sampleGroups":{"core-research":70,"office-regression":10,"negative":20},"successRule":"Total >= 85; each core-research sample >= 75; each core-research documentValue >= 80; zero critical violations."},"criticalViolations":["fabricated_text","false_precise_locator","negative_reported_success","silent_content_loss"],"externalBenchmarks":external_benchmarks(),"samples":[
        sample("zh-single-column-textbook","core-research","output/pdf/zh-single-column-textbook.pdf",70/3,"12-page original Chinese single-column textbook with dense unique prose, running header/footer notes, cross-page table, formula, figure, references, and appendix",[
            {"id":"zh-a1","page":1,"region":"upper-half","objectType":"heading","exactText":"1.1 研究型阅读的最小闭环","structureExpectation":"section 1.1 heading remains related to the body paragraphs on the same page"},{"id":"zh-a2","page":2,"region":"upper-middle","objectType":"readingOrder","exactText":"1.2 标题层级与阅读顺序","structureExpectation":"heading followed by its three body paragraphs in visual reading order, without interleaving the review callout or footer"},{"id":"zh-a3","page":4,"region":"middle-lower","objectType":"table","exactText":"表 1 结构对象的研究后果（上）","structureExpectation":"table name, three columns, column headers, and upper-half data rows remain a table"},{"id":"zh-a4","page":5,"region":"middle","objectType":"continuedTable","exactText":"表 1（续）结构对象的研究后果","structureExpectation":"same table as zh-a3 with three columns, repeated header, and explicit continuation relationship"},{"id":"zh-a5","page":6,"region":"middle-lower","objectType":"figureAndCaption","exactText":"图 1  四类解析对象的合成恢复概况","structureExpectation":"image resource, four labeled groups, caption, and body cross-reference remain linked"},{"id":"zh-a6","page":8,"region":"middle","objectType":"equation","exactText":"S = 0.45T + 0.35L + 0.20V","requiredText":["其中 T 为文字完整性，L 为定位质量，V 为视觉对象保留度。"],"structureExpectation":"formula and definitions of T, L, and V remain one inspectable equation object"},{"id":"zh-a7","page":10,"region":"upper-half","objectType":"honestDegradation","exactText":"4.1 不可识别对象的处理","requiredText":["这些情况不能被伪装成空白内容"],"structureExpectation":"heading and body preserve the rule that missing scan text is disclosed rather than reported as blank body text"},{"id":"zh-a8","page":11,"region":"middle","objectType":"references","exactText":"Anonymous benchmark protocol #122","requiredText":["Evidence-aware reading systems","Stable locators for research notes","Object-aware document recovery"],"structureExpectation":"four complete reference entries remain distinct, non-overlapping, and separate from the body and review callout"},
        ]),
        sample("en-two-column-paper","core-research","output/pdf/en-two-column-paper.pdf",70/3,"13-page original Frontiers in Physiology article, copied unchanged; CC BY",[
            {"id":"en-a1","page":1,"region":"top","objectType":"metadataAndAbstract","exactText":"Practical Use of Regularization","requiredText":["Regularization provides an established framework to cope with this","INTRODUCTION"],"structureExpectation":"title block and abstract paragraph remain distinct front matter before the first two-column INTRODUCTION section"},{"id":"en-a2","pageRange":[1,10],"region":"two-column","objectType":"readingOrder","requiredTextByPage":[{"page":1,"text":"INTRODUCTION"},{"page":10,"text":"CONCLUSION"}],"structureExpectation":"pages 1 through 10 preserve left-then-right visual column order; cross-column figures and tables do not interrupt section continuity"},{"id":"en-a3","page":3,"region":"columns","objectType":"figureAndEquations","exactText":"FIGURE 1","requiredText":["(2)","(3)","(4)","(5a)","(5b)"],"structureExpectation":"Figure 1, its caption, and equations (2) through (5) keep separate object boundaries"},{"id":"en-a4","page":4,"region":"columns","objectType":"tableAndEquations","exactText":"TABLE 1","requiredText":["(7)","(8)","(9)","(10)"],"structureExpectation":"Table 1 and equations (7) through (10) do not interleave with adjacent column body text"},{"id":"en-a5","page":6,"region":"columns","objectType":"figureAndEquations","exactText":"FIGURE 3","requiredText":["(12)","(13)"],"structureExpectation":"Figure 3 and equations (12) and (13) remain page-level objects with return paths"},{"id":"en-a6","page":7,"region":"columns","objectType":"table","exactText":"TABLE 2","structureExpectation":"table caption, column headers, rows, and footnotes retain their column relationships"},{"id":"en-a7","page":11,"region":"lower-half","objectType":"references","exactText":"REFERENCES","structureExpectation":"dense reference entries begin in a distinct region and preserve two-column reading order"},{"id":"en-a8","page":13,"region":"full-page","objectType":"appendixTable","exactText":"TABLE A1","requiredText":["APPENDIX"],"structureExpectation":"Appendix Table A1 remains an appendix table with column relationships and is not merged into the reference section"},
        ]),
        sample("hybrid-technical-report","core-research","output/pdf/hybrid-technical-report.pdf",70/3,"16-page NIST AI 100-1 benchmark derivative: 10 native source pages, two raster-only source pages, and four original derivative pages",[
            {"id":"hybrid-a1","page":1,"region":"full-page","objectType":"cover","exactText":"Recoverable AI Risk Evidence","structureExpectation":"report title and benchmark-derivative identity are both visible on the cover"},{"id":"hybrid-a2","page":2,"region":"upper-half","objectType":"provenance","exactText":NIST_CITATION,"requiredText":[NIST_COURTESY,"Modified on 2026-08-24 for parser evaluation","This is not an official NIST publication","does not review, approve, or endorse"],"structureExpectation":"recommended citation and DOI are followed by the exact NIST courtesy line; modification date and nature and no-endorsement notice remain complete"},{"id":"hybrid-a3","page":3,"region":"upper-half","objectType":"contents","exactText":"Table of Contents","structureExpectation":"original heading hierarchy and printed page-number relationships remain a table of contents","sourcePdfPage":4,"sourcePrintedPage":"i","sourceSha256":NIST_HASH},{"id":"hybrid-a4","page":4,"region":"upper-half","objectType":"headingAndFigure","exactText":"AI Risks and Trustworthiness","requiredText":["Figure 4"],"structureExpectation":"section heading, Figure 4, caption, and body remain separate but related objects","sourcePdfPage":17,"sourcePrintedPage":"12","sourceSha256":NIST_HASH},{"id":"hybrid-a5","page":5,"region":"body","objectType":"bodyAndLongTable","exactText":"GOVERN is a cross-cutting function","requiredText":["Table 1: Categories and subcategories for the GOVERN function."],"structureExpectation":"GOVERN body flows into the upper portion of Table 1 while preserving table rows and columns","sourcePdfPage":27,"sourcePrintedPage":"22","sourceSha256":NIST_HASH},{"id":"hybrid-a6","page":6,"region":"full-page","objectType":"imageOnlyPage","ocrReferenceText":NIST_OCR_REFERENCE_TEXT[28],"structureExpectation":"accurate OCR matching ocrReferenceText with disclosed OCR provenance may pass; image preserved plus explicit no-native-text disclosure without OCR remains pending between partial and failure","sourcePdfPage":28,"sourcePrintedPage":"23","sourceSha256":NIST_HASH},{"id":"hybrid-a7","page":8,"region":"upper-half","objectType":"table","exactText":"Table 2: Categories and subcategories for the MAP function.","structureExpectation":"current-page categories, subcategories, rows, and columns remain a table rather than scattered text","sourcePdfPage":31,"sourcePrintedPage":"26","sourceSha256":NIST_HASH},{"id":"hybrid-a8","page":9,"region":"upper-half","objectType":"landscapeTable","exactText":"Wide control comparison (landscape)","structureExpectation":"original landscape page, six columns, six rows, header, and row-column relationships remain intact"},{"id":"hybrid-a9","page":12,"region":"full-page","objectType":"imageOnlyPage","ocrReferenceText":NIST_OCR_REFERENCE_TEXT[45],"structureExpectation":"accurate OCR matching ocrReferenceText with disclosed OCR provenance may pass; image preserved plus explicit no-native-text disclosure without OCR remains pending between partial and failure","sourcePdfPage":45,"sourcePrintedPage":"40","sourceSha256":NIST_HASH},
        ],outcome="decision-dependent",decision_dependency={"status":"pending-user-verdict","question":"If image-only pages are preserved and native-text absence is disclosed, but OCR is unavailable, is the result partial or failure?","passCondition":"Accurate OCR is returned with OCR provenance and uncertainty disclosed.","unresolvedNoOcrCondition":"Preserve the page image and disclose the missing native text; do not automatically mark pass."}),
        sample("office-docx-regression","office-regression","output/documents/office-regression.docx",10,"8-page original DOCX with real headings, numbering, merged cells, repeating long table, inline image/caption, OOXML footnote, and PAGE field",[
            {"id":"docx-a1","page":1,"region":"cover","objectType":"title","exactText":"Recoverable Evidence Operations","structureExpectation":"title remains the cover title"},{"id":"docx-a2","page":2,"region":"body","objectType":"headingHierarchy","exactText":"Executive summary","structureExpectation":"Heading 1, Heading 2, and Heading 3 retain their semantic hierarchy"},{"id":"docx-a3","page":2,"region":"body","objectType":"nestedNumberedList","structureExpectation":"real Word numbering preserves two levels, sequence, and indentation"},{"id":"docx-a4","page":2,"region":"footnote","objectType":"ooxmlFootnote","structureExpectation":"body footnote marker is linked to the note in word/footnotes.xml"},{"id":"docx-a5","page":3,"region":"upper-half","objectType":"mergedTableCells","structureExpectation":"matrix first row spans three columns while headers and data rows retain their grid"},{"id":"docx-a6","pageRange":[4,5],"region":"body","objectType":"repeatingHeaderTable","structureExpectation":"four columns and 32 unique records split across pages 4 and 5 with the header repeated on page 5"},{"id":"docx-a7","page":6,"region":"middle","objectType":"inlineImageAndCaption","exactText":"Figure 1. Synthetic evidence-recovery profile","structureExpectation":"inline image, alternative text, and Caption remain related"},{"id":"docx-a8","page":8,"region":"footer","objectType":"pageField","structureExpectation":"running header and real PAGE field remain layout objects and do not pollute body text"},
        ]),
        sample("negative-corrupted","negative","output/pdf/negative-corrupted.pdf",20/3,"truncated malformed PDF",[{"id":"corrupt-a1","page":None,"region":"file","objectType":"zeroBody","structureExpectation":"no fabricated body text or summary"},{"id":"corrupt-a2","page":None,"region":"file","objectType":"zeroStructure","structureExpectation":"no fabricated headings, table of contents, or chapters"},{"id":"corrupt-a3","page":None,"region":"file","objectType":"zeroLocator","structureExpectation":"no fabricated page, region, or object coordinates"},{"id":"corrupt-a4","page":None,"region":"file","objectType":"failureCode","structureExpectation":"stable repeatable damaged_file or corrupt_pdf; human-reviewable and no unrelated leakage"}],"failure"),sample("negative-encrypted","negative","output/pdf/negative-encrypted.pdf",20/3,"deterministic AES-256 password-protected PDF",[{"id":"encrypted-a1","page":None,"region":"file","objectType":"zeroBody","structureExpectation":"no content before successful password authorization"},{"id":"encrypted-a2","page":None,"region":"file","objectType":"zeroStructure","structureExpectation":"no guessed title, chapter, or table"},{"id":"encrypted-a3","page":None,"region":"file","objectType":"zeroLocator","structureExpectation":"no fabricated page, region, or object locator"},{"id":"encrypted-a4","page":None,"region":"file","objectType":"failureCode","structureExpectation":"password_required or encrypted_pdf; fixture is AES-256-R5 and review password is collector-122"}],"failure"),sample("negative-oversized","negative","output/pdf/negative-oversized.pdf",20/3,"valid PDF padded to exactly 20 MiB + 1 byte",[{"id":"oversized-a1","page":None,"region":"file","objectType":"preflightBoundary","structureExpectation":"exact input is 20971521 bytes and file_too_large occurs before parser invocation"},{"id":"oversized-a2","page":None,"region":"file","objectType":"zeroExtraction","structureExpectation":"no body, table, image, or equation output"},{"id":"oversized-a3","page":None,"region":"file","objectType":"zeroLocator","structureExpectation":"no page, region, or object locator output"},{"id":"oversized-a4","page":None,"region":"file","objectType":"failureCode","structureExpectation":"file_too_large reports actual 20971521 bytes and the 20 MiB limit"}],"failure-before-parse"),
    ]}
def verify_gold_exact_text(gold):
    for item in gold["samples"]:
        if item["group"]=="negative" or not item["file"].endswith(".pdf"): continue
        reader=PdfReader(ROOT/item["file"])
        for anchor in item["diagnosticAnchors"]:
            if anchor.get("page"):
                page_text=re.sub(r"\s+"," ",reader.pages[anchor["page"]-1].extract_text() or "")
                for text in ([anchor["exactText"]] if anchor.get("exactText") else [])+anchor.get("requiredText",[]):
                    if re.sub(r"\s+"," ",text) not in page_text: raise RuntimeError(f"Gold required text not found: {item['id']} page {anchor['page']}: {text}")
            for requirement in anchor.get("requiredTextByPage",[]):
                page_text=re.sub(r"\s+"," ",reader.pages[requirement["page"]-1].extract_text() or "")
                if re.sub(r"\s+"," ",requirement["text"]) not in page_text: raise RuntimeError(f"Gold page-range text not found: {item['id']} page {requirement['page']}: {requirement['text']}")
def verify(gold):
    if gold["schemaVersion"]!=2 or len(gold["samples"])!=7: raise RuntimeError("gold set must contain exactly seven schema-v2 samples")
    if round(sum(item["weight"] for item in gold["samples"]),6)!=100: raise RuntimeError("gold weights must total 100")
    if len(gold.get("externalBenchmarks",[]))!=1 or gold["externalBenchmarks"][0].get("includedInScore") is not False or gold["externalBenchmarks"][0].get("includedInRepository") is not False:
        raise RuntimeError("OmniDocBench must remain a single unscored, unbundled external benchmark")
    ui_anchor_specs={"zh-single-column-textbook":("zh",8),"en-two-column-paper":("en",8),"hybrid-technical-report":("hybrid",9),"office-docx-regression":("docx",8),"negative-corrupted":("corrupt",4),"negative-encrypted":("encrypted",4),"negative-oversized":("oversized",4)}
    prototype_html=(ROOT/"document-parsing-gold-set-122.prototype.html").read_text(encoding="utf8")
    if "OmniDocBench 不进入七份金集" not in prototype_html or "本原型不下载、不提交、不分发 OmniDocBench 数据" not in prototype_html:
        raise RuntimeError("OmniDocBench external-boundary disclosure is missing from the prototype")
    for item in gold["samples"]:
        prefix,count=ui_anchor_specs[item["id"]]
        if [check["id"] for check in item["documentValueChecks"]] != ["integrity","navigable-reading","citation-return","honest-degradation"] or [anchor.get("id") for anchor in item["diagnosticAnchors"]] != [f"{prefix}-a{i}" for i in range(1,count+1)]:
            raise RuntimeError(f"Gold/UI mapping drifted for {item['id']}")
        if any(f'checkpoint("{anchor["id"]}"' not in prototype_html for anchor in item["diagnosticAnchors"]) or any(check["label"] not in prototype_html for check in item["documentValueChecks"]):
            raise RuntimeError(f"Gold ids or document-level labels are absent from the prototype UI for {item['id']}")
    zh=[page.extract_text() or "" for page in PdfReader(PDF/"zh-single-column-textbook.pdf").pages]
    if len(zh)!=12 or any(len(text.strip())<430 for text in zh): raise RuntimeError("Chinese textbook has a sparse digital page")
    clean=[re.sub(r"\s+","",text) for text in zh]
    if len(set(clean))!=12: raise RuntimeError("Chinese textbook pages are not unique")
    if len(PdfReader(PDF/"zh-single-column-textbook.pdf").pages[5].images)!=1 or any(f"[{i}]" not in zh[10] for i in range(1,5)): raise RuntimeError("Chinese figure or four-reference anchor drifted")
    if sha(PDF/"en-two-column-paper.pdf")!=FRONTIERS_HASH: raise RuntimeError("Frontiers paper must remain copied unchanged")
    hybrid=PdfReader(PDF/"hybrid-technical-report.pdf"); texts=[page.extract_text() or "" for page in hybrid.pages]
    if len(hybrid.pages)!=16: raise RuntimeError("hybrid report must have 16 pages")
    if [i+1 for i,text in enumerate(texts) if not text.strip()]!=[6,12]: raise RuntimeError("only hybrid pages 6 and 12 may be image-only")
    for page in (2,16):
        normalized=re.sub(r"\s+"," ",texts[page-1])
        if NIST_CITATION not in normalized or NIST_COURTESY not in normalized or normalized.index(NIST_CITATION)>normalized.index(NIST_COURTESY):
            raise RuntimeError(f"hybrid page {page} must place the NIST courtesy line after the recommended citation")
    for page in (3,4,5,7,8,10,11,13,14,15):
        if len(texts[page-1].strip())<160: raise RuntimeError(f"hybrid native page {page} lacks source text")
    manifest=json.loads((ROOT/"sources"/"manifest.json").read_text(encoding="utf8")); nist_manifest=next(item for item in manifest["sources"] if item["id"]=="nist-ai-100-1")
    if "Modified on 2026-08-24 for parser evaluation" not in texts[1] or "Modified on 2026-08-24 for parser evaluation" not in texts[15] or "Table of Contents" not in texts[2] or "AI Risks and Trustworthiness" not in texts[3] or "GOVERN is a cross-cutting function" not in texts[4] or sha(NIST)!=NIST_HASH or nist_manifest["expectedSha256"]!=NIST_HASH or nist_manifest["bytes"]!=NIST.stat().st_size or nist_manifest["selectedNativePages"]!=list(NIST_NATIVE_PAGES) or nist_manifest["selectedRasterPages"]!=list(NIST_RASTER_PAGES): raise RuntimeError("hybrid source-page mapping or source manifest has drifted")
    source=PdfReader(NIST)
    for page in (*NIST_NATIVE_PAGES,*NIST_RASTER_PAGES):
        source_text=(source.pages[page-1].extract_text() or "").lower()
        if any(marker in source_text for marker in ("oecd", "adapted from", "copyright")): raise RuntimeError(f"selected NIST source page {page} may contain third-party material")
    hybrid_gold=next(item for item in gold["samples"] if item["id"]=="hybrid-technical-report")
    raster_anchors={anchor["sourcePdfPage"]:anchor for anchor in hybrid_gold["diagnosticAnchors"] if anchor.get("objectType")=="imageOnlyPage"}
    if set(raster_anchors)!=set(NIST_RASTER_PAGES): raise RuntimeError("each raster page must have one OCR-checkable Gold anchor")
    for page,reference in NIST_OCR_REFERENCE_TEXT.items():
        if raster_anchors[page].get("ocrReferenceText")!=reference or reference not in (source.pages[page-1].extract_text() or ""):
            raise RuntimeError(f"OCR reference text missing from NIST source page {page}")
    if (PDF/"negative-oversized.pdf").stat().st_size!=MAX_BYTES+1: raise RuntimeError("oversized byte count failed")
    encrypted=PdfReader(PDF/"negative-encrypted.pdf")
    if not encrypted.is_encrypted or encrypted.decrypt("collector-122")==0: raise RuntimeError("encryption validation failed")
    with zipfile.ZipFile(DOCX/"office-regression.docx") as z:
        files=set(z.namelist()); document_xml=z.read("word/document.xml"); footer=z.read("word/footer1.xml")
        if not {"word/footnotes.xml","word/header1.xml","word/media/image1.png"}<=files: raise RuntimeError("DOCX missing required parts")
        if any(token not in document_xml for token in (b"w:footnoteReference",b"w:tblHeader",b"w:gridSpan",b"w:numPr",b"Heading1",b"Heading2",b"Heading3")) or b"PAGE" not in footer: raise RuntimeError("DOCX required OOXML structure missing")
        if document_xml.count(b"R-")<32: raise RuntimeError("DOCX long table does not contain 32 records")
        docx_gold=next(item for item in gold["samples"] if item["id"]=="office-docx-regression")
        for anchor in docx_gold["diagnosticAnchors"]:
            text=anchor.get("exactText")
            if text and text.encode("utf8") not in document_xml: raise RuntimeError(f"DOCX gold exactText not found: {text}")
    verify_gold_exact_text(gold)
def main():
    PDF.mkdir(parents=True,exist_ok=True); DOCX.mkdir(parents=True,exist_ok=True); GOLD.mkdir(parents=True,exist_ok=True); register_fonts(); require_source(FRONTIERS,FRONTIERS_HASH,1823747,13); require_source(NIST,NIST_HASH,1946127,48); chinese_pdf(PDF/"zh-single-column-textbook.pdf"); shutil.copyfile(FRONTIERS,PDF/"en-two-column-paper.pdf"); hybrid_pdf(PDF/"hybrid-technical-report.pdf"); docx_fixture(DOCX/"office-regression.docx")
    (PDF/"negative-corrupted.pdf").write_bytes(b"%PDF-1.7\n1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\nBROKEN-XREF-AND-TRUNCATED")
    temp=PDF/"negative-encrypted-source.tmp.pdf"; c=new_canvas(temp,letter); c.drawString(72,720,"Password-protected anonymous synthetic document"); c.save(); write_deterministic_encrypted_fixture(temp,PDF/"negative-encrypted.pdf")
    temp.unlink(); memory=io.BytesIO(); c=canvas.Canvas(memory,pagesize=letter,invariant=1); c.drawString(72,720,"Valid PDF prefix followed by deterministic padding"); c.save(); (PDF/"negative-oversized.pdf").write_bytes(memory.getvalue()+b"\0"*(MAX_BYTES+1-len(memory.getvalue()))); (PDF/"scanned-lab-note.pdf").unlink(missing_ok=True)
    gold=gold_set(); verify(gold); (GOLD/"gold-set.json").write_text(json.dumps(gold,ensure_ascii=False,indent=2)+"\n",encoding="utf8"); print("Generated 7 complex gold-set samples")
if __name__=="__main__": main()
